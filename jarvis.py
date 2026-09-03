# -*- coding: utf-8 -*-
import os
import re
import io
import json
import time
import shutil
import zipfile
import hashlib
import subprocess
import unicodedata
from pathlib import Path
from urllib.parse import quote_plus, urlparse, unquote
from urllib.request import Request, urlopen

import pyperclip
from ddgs import DDGS
from send2trash import send2trash
from pywinauto import Desktop, keyboard
from docx import Document


# ============================================================
# AYARLAR
# ============================================================

KULLANICI = Path.home()
MASAUSTU = KULLANICI / "Desktop"
BELGELER = KULLANICI / "Documents"
INDIRMELER = KULLANICI / "Downloads"
JARVIS_KLASORU = KULLANICI / "Jarvis"

ONEDRIVE = Path(os.environ["OneDrive"]) if os.environ.get("OneDrive") else None

MAKSIMUM_BELGE = 3
MAKSIMUM_DOSYA_BOYUTU = 25 * 1024 * 1024
BELGE_UZANTILARI = {".pdf", ".docx", ".doc"}
ARAMA_MOTORLARI = ["duckduckgo", "brave", "yahoo"]

CLAUDE_APP_ID = "Claude_pzs8sxrjxfjjc!Claude"

PROGRAMLAR = {
    "not defteri": {
        "ac": ["notepad.exe"],
        "process": "notepad",
    },
    "word": {
        "ac": ["cmd", "/c", "start", "", "winword"],
        "process": "WINWORD",
    },
    "powerpoint": {
        "ac": ["cmd", "/c", "start", "", "powerpnt"],
        "process": "POWERPNT",
    },
    "chrome": {
        "ac": ["cmd", "/c", "start", "", "chrome"],
        "process": "chrome",
    },
    "claude": {
        "ac": ["explorer.exe", f"shell:AppsFolder\\{CLAUDE_APP_ID}"],
        "process": "Claude",
    },
}

PROGRAM_ESANLAMLILAR = {
    "not defteri": "not defteri",
    "notepad": "not defteri",
    "word": "word",
    "microsoft word": "word",
    "powerpoint": "powerpoint",
    "power point": "powerpoint",
    "chrome": "chrome",
    "google chrome": "chrome",
    "claude": "claude",
    "claude desktop": "claude",
}


# ============================================================
# GENEL YARDIMCILAR
# ============================================================


def konus(metin):
    """Jarvis'i ElevenLabs OAuth + Alper sesiyle konusturur."""

    try:
        import ctypes

        metin = str(metin).strip()

        if not metin:
            return

        voice_id = "HllA1j2zLOqUQ4kLjMmK"
        model_id = "eleven_multilingual_v2"

        # Ayni cumle tekrar soylenirse yeniden ElevenLabs
        # kullanimi yapmamak icin yerel ses onbellegi.
        cache_klasoru = JARVIS_KLASORU / "ses_cache"
        cache_klasoru.mkdir(
            parents=True,
            exist_ok=True
        )

        kimlik = hashlib.sha256(
            (
                voice_id
                + "|"
                + model_id
                + "|"
                + metin
            ).encode("utf-8")
        ).hexdigest()

        ses_dosyasi = (
            cache_klasoru
            / f"{kimlik}.mp3"
        )

        if not ses_dosyasi.exists():

            komut = [
                "cmd",
                "/c",
                "elevenlabs",
                "text-to-speech",
                "convert",
                "--model-id",
                model_id,
                "--voice-id",
                voice_id,
                "--text",
                metin,
                "--output",
                str(ses_dosyasi)
            ]

            sonuc = subprocess.run(
                komut,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=120,
                creationflags=getattr(
                    subprocess,
                    "CREATE_NO_WINDOW",
                    0
                )
            )

            if sonuc.returncode != 0:
                hata = (
                    sonuc.stderr.strip()
                    or sonuc.stdout.strip()
                    or "Bilinmeyen ElevenLabs hatasi"
                )

                print(
                    "Jarvis: ElevenLabs OAuth ses hatasi:",
                    hata
                )

                return

            if (
                not ses_dosyasi.exists()
                or ses_dosyasi.stat().st_size == 0
            ):
                print(
                    "Jarvis: ElevenLabs ses dosyasi olusturmadi."
                )
                return

        # MP3 dosyasini Windows MCI ile oynat.
        mci = ctypes.windll.winmm.mciSendStringW

        alias = (
            "jarvis_ses_"
            + str(os.getpid())
        )

        mci(
            f"close {alias}",
            None,
            0,
            None
        )

        sonuc = mci(
            f'open "{ses_dosyasi}" type mpegvideo alias {alias}',
            None,
            0,
            None
        )

        if sonuc != 0:
            print(
                "Jarvis: Ses dosyasi acilamadi."
            )
            return

        try:
            mci(
                f"play {alias} wait",
                None,
                0,
                None
            )

        finally:
            mci(
                f"close {alias}",
                None,
                0,
                None
            )

    except subprocess.TimeoutExpired:
        print(
            "Jarvis: ElevenLabs ses istegi zaman asimina ugradi."
        )

    except Exception as e:
        print(
            "Jarvis: Ses sistemi hatasi:",
            e
        )


def norm(metin):
    metin = str(metin).casefold().replace("\u0307", "")
    metin = metin.translate(str.maketrans({
        "ı": "i",
        "ş": "s",
        "ğ": "g",
        "ü": "u",
        "ö": "o",
        "ç": "c",
    }))
    metin = unicodedata.normalize("NFKD", metin)
    metin = "".join(c for c in metin if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", metin).strip()


def guvenli_ad(ad):
    ad = re.sub(r'[<>:"/\\|?*]+', "_", str(ad)).strip(" .")
    return ad or "dosya"


def benzersiz_yol(yol):
    yol = Path(yol)
    if not yol.exists():
        return yol

    i = 2
    while True:
        aday = yol.with_name(f"{yol.stem}_{i}{yol.suffix}")
        if not aday.exists():
            return aday
        i += 1


def arama_kokleri():
    adaylar = [MASAUSTU, BELGELER, INDIRMELER, JARVIS_KLASORU]

    if ONEDRIVE:
        adaylar.extend([
            ONEDRIVE / "Desktop",
            ONEDRIVE / "Documents",
        ])

    sonuc = []
    gorulen = set()

    for p in adaylar:
        if not p.exists():
            continue

        try:
            anahtar = norm(str(p.resolve()))
        except Exception:
            anahtar = norm(str(p))

        if anahtar not in gorulen:
            gorulen.add(anahtar)
            sonuc.append(p)

    return sonuc


def tarayicida_ac(url):
    try:
        os.startfile(url)
        return True
    except Exception as e:
        print("Jarvis: Tarayıcı açılamadı:", e)
        return False


def klasoru_ac(yol):
    try:
        os.startfile(str(yol))
        return True
    except Exception as e:
        print("Jarvis: Klasör açılamadı:", e)
        return False


def dosyayi_ac(yol):
    try:
        os.startfile(str(yol))
        return True
    except Exception as e:
        print("Jarvis: Dosya açılamadı:", e)
        return False


# ============================================================
# PROGRAM AÇ / KAPAT
# ============================================================

def program_adi_bul(metin):
    n = norm(metin)
    for takma, gercek in PROGRAM_ESANLAMLILAR.items():
        if norm(takma) in n:
            return gercek
    return None


def program_ac(ad):
    bilgi = PROGRAMLAR.get(ad)

    if not bilgi:
        print("Jarvis: Bu programı bilmiyorum.")
        return

    try:
        subprocess.Popen(bilgi["ac"])
        print(f"Jarvis: {ad} açılıyor.")
    except Exception as e:
        print(f"Jarvis: {ad} açılamadı:", e)


def program_kapat(ad):
    bilgi = PROGRAMLAR.get(ad)

    if not bilgi:
        print("Jarvis: Bu programı bilmiyorum.")
        return

    process = bilgi["process"]

    ps = (
        f"$p=Get-Process -Name '{process}' -ErrorAction SilentlyContinue; "
        "if($p){$p | ForEach-Object { $_.CloseMainWindow() | Out-Null }; exit 0} "
        "else {exit 1}"
    )

    try:
        r = subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps],
            capture_output=True,
            text=True
        )

        if r.returncode == 0:
            print(f"Jarvis: {ad} için kapatma isteği gönderildi.")
        else:
            print(f"Jarvis: {ad} açık görünmüyor.")

    except Exception as e:
        print(f"Jarvis: {ad} kapatılamadı:", e)


# ============================================================
# DOSYA / KLASÖR ARAMA
# ============================================================

def dosya_puani(path, sorgu):
    q = norm(sorgu)
    ad = norm(path.name)
    govde = norm(path.stem)

    if q == ad:
        return 100

    if q == govde:
        return 95

    if ad.startswith(q) or govde.startswith(q):
        return 80

    if q in ad or q in govde:
        return 60

    kelimeler = [x for x in q.split() if len(x) >= 2]

    if kelimeler and all(x in ad for x in kelimeler):
        return 50

    return 0


def dosya_bul(sorgu, kok=None, en_fazla=30):
    sorgu = str(sorgu).strip().strip('"')

    dogrudan = Path(os.path.expandvars(os.path.expanduser(sorgu)))

    if dogrudan.exists() and dogrudan.is_file():
        return [dogrudan]

    kokler = [Path(kok)] if kok else arama_kokleri()
    bulunan = []

    for root in kokler:
        try:
            for dirpath, _, filenames in os.walk(root):
                for filename in filenames:
                    p = Path(dirpath) / filename
                    puan = dosya_puani(p, sorgu)

                    if puan:
                        bulunan.append((puan, p))
        except Exception:
            continue

    bulunan.sort(key=lambda x: (-x[0], len(str(x[1]))))

    sonuc = []
    gorulen = set()

    for _, p in bulunan:
        anahtar = norm(str(p))

        if anahtar not in gorulen:
            gorulen.add(anahtar)
            sonuc.append(p)

        if len(sonuc) >= en_fazla:
            break

    return sonuc


def klasor_bul(sorgu, kok=None, en_fazla=30):
    sorgu = str(sorgu).strip().strip('"')

    dogrudan = Path(os.path.expandvars(os.path.expanduser(sorgu)))

    if dogrudan.exists() and dogrudan.is_dir():
        return [dogrudan]

    q = norm(sorgu)
    kokler = [Path(kok)] if kok else arama_kokleri()
    bulunan = []

    for root in kokler:
        try:
            for dirpath, dirnames, _ in os.walk(root):
                for dirname in dirnames:
                    p = Path(dirpath) / dirname
                    ad = norm(dirname)

                    if q == ad:
                        puan = 100
                    elif ad.startswith(q):
                        puan = 80
                    elif q in ad:
                        puan = 60
                    else:
                        continue

                    bulunan.append((puan, p))

        except Exception:
            continue

    bulunan.sort(key=lambda x: (-x[0], len(str(x[1]))))

    sonuc = []
    gorulen = set()

    for _, p in bulunan:
        anahtar = norm(str(p))

        if anahtar not in gorulen:
            gorulen.add(anahtar)
            sonuc.append(p)

        if len(sonuc) >= en_fazla:
            break

    return sonuc


def tek_dosya(sorgu, kok=None):
    bulunan = dosya_bul(sorgu, kok=kok)

    if not bulunan:
        print(f"Jarvis: '{sorgu}' adlı dosyayı bulamadım.")
        return None

    en_iyi = dosya_puani(bulunan[0], sorgu)

    ayni = [
        p for p in bulunan
        if dosya_puani(p, sorgu) == en_iyi
    ]

    if len(ayni) > 1:
        print("Jarvis: Birden fazla eşleşme buldum. İşlem yapmıyorum:")

        for p in ayni[:10]:
            print(" -", p)

        return None

    return bulunan[0]


def tek_klasor(sorgu, kok=None):
    bulunan = klasor_bul(sorgu, kok=kok)

    if not bulunan:
        print(f"Jarvis: '{sorgu}' adlı klasörü bulamadım.")
        return None

    if len(bulunan) > 1 and norm(bulunan[0].name) == norm(bulunan[1].name):
        print("Jarvis: Aynı isimde birden fazla klasör buldum:")

        for p in bulunan[:10]:
            print(" -", p)

        return None

    return bulunan[0]


# ============================================================
# DOSYA İŞLEMLERİ
# ============================================================


def dogal_ac_komutu(sorgu):
    """
    'kisisel bilgi kayit formu ac'
    gibi dogal komutlarda dosya veya klasoru bulup acar.
    Belirsizlik varsa kendi kendine karar vermez.
    """

    q = norm(sorgu).strip()

    if not q:
        return False

    # Konusmada kullanilabilecek gereksiz kelimeleri temizle.
    q = re.sub(
        r"\s+(?:dosyami|dosyayi|dosyasini|dosya|"
        r"klasorumu|klasoru|klasorunu|klasor)$",
        "",
        q
    ).strip()

    if not q:
        return False

    dosyalar = dosya_bul(
        q,
        en_fazla=10
    )

    klasorler = klasor_bul(
        q,
        en_fazla=10
    )

    # Sadece dosya bulundu.
    if dosyalar and not klasorler:
        en_iyi = dosya_puani(
            dosyalar[0],
            q
        )

        ayni = [
            p for p in dosyalar
            if dosya_puani(p, q) == en_iyi
        ]

        if len(ayni) > 1:
            print(
                "Jarvis: Birden fazla dosya eslesmesi buldum. "
                "Hangisini istedigini belirt:"
            )

            for p in ayni[:10]:
                print(" -", p)

            return True

        dosyayi_ac(
            dosyalar[0]
        )

        print(
            "Jarvis: Acildi:",
            dosyalar[0]
        )

        return True

    # Sadece klasor bulundu.
    if klasorler and not dosyalar:
        if (
            len(klasorler) > 1
            and norm(klasorler[0].name)
            == norm(klasorler[1].name)
        ):
            print(
                "Jarvis: Ayni isimde birden fazla klasor buldum:"
            )

            for p in klasorler[:10]:
                print(" -", p)

            return True

        klasoru_ac(
            klasorler[0]
        )

        print(
            "Jarvis: Acildi:",
            klasorler[0]
        )

        return True

    # Hem dosya hem klasor bulunduysa otomatik tahmin yapma.
    if dosyalar and klasorler:
        print(
            "Jarvis: Hem dosya hem klasor eslesmesi buldum. "
            "Yanlis bir sey acmamak icin secim yapmiyorum."
        )

        print("Dosya:")
        for p in dosyalar[:3]:
            print(" -", p)

        print("Klasor:")
        for p in klasorler[:3]:
            print(" -", p)

        return True

    print(
        f"Jarvis: '{sorgu}' icin dosya veya klasor bulamadim."
    )

    return True



def dosya_ac_komutu(sorgu):
    p = tek_dosya(sorgu)

    if p:
        dosyayi_ac(p)
        print("Jarvis: Açıldı:", p)


def klasor_ac_komutu(sorgu):
    p = tek_klasor(sorgu)

    if p:
        klasoru_ac(p)
        print("Jarvis: Açıldı:", p)


def dosya_sil_komutu(sorgu):
    p = tek_dosya(sorgu)

    if not p:
        return

    print("Jarvis: Silinecek dosya:", p)

    onay = norm(
        input(
            "Jarvis: Geri Dönüşüm Kutusu'na göndereyim mi? "
            "(EVET/HAYIR): "
        )
    )

    if onay != "evet":
        print("Jarvis: İptal edildi.")
        return

    try:
        send2trash(str(p))
        print("Jarvis: Dosya Geri Dönüşüm Kutusu'na gönderildi.")
    except Exception as e:
        print("Jarvis: Silinemedi:", e)


def dosya_yeniden_adlandir(sorgu, yeni_ad, kok=None):
    p = tek_dosya(sorgu, kok=kok)

    if not p:
        return

    yeni_ad = yeni_ad.strip().strip('"')

    if not Path(yeni_ad).suffix:
        yeni_ad += p.suffix

    hedef = p.with_name(guvenli_ad(yeni_ad))

    if hedef.exists():
        print("Jarvis: Bu isimde bir dosya zaten var. Üzerine yazmıyorum.")
        return

    try:
        p.rename(hedef)
        print("Jarvis: Yeni ad:", hedef.name)
    except Exception as e:
        print("Jarvis: Yeniden adlandırılamadı:", e)


def dosya_tasi_veya_kopyala(sorgu, hedef_klasor_adi, kopyala=False):
    p = tek_dosya(sorgu)

    if not p:
        return

    hedef_klasor = tek_klasor(hedef_klasor_adi)

    if not hedef_klasor:
        return

    hedef = hedef_klasor / p.name

    if hedef.exists():
        print("Jarvis: Hedefte aynı isimde dosya var. Üzerine yazmıyorum.")
        return

    try:
        if kopyala:
            shutil.copy2(p, hedef)
            print("Jarvis: Dosya kopyalandı:", hedef)
        else:
            shutil.move(str(p), str(hedef))
            print("Jarvis: Dosya taşındı:", hedef)

    except Exception as e:
        print("Jarvis: İşlem başarısız:", e)


# ============================================================
# GOOGLE / HIZLI ARAŞTIRMA
# ============================================================

def google_ara(sorgu):
    tarayicida_ac(
        "https://www.google.com/search?q=" + quote_plus(sorgu)
    )

    print("Jarvis: Google araması açıldı.")


def ddgs_yedekli_ara(sorgu, adet=3):
    son_hata = None

    for motor in ARAMA_MOTORLARI:
        try:
            sonuc = list(
                DDGS().text(
                    sorgu,
                    max_results=max(adet * 3, 10),
                    backend=motor
                )
            )

            if sonuc:
                return sonuc[:adet]

        except Exception as e:
            son_hata = e

    if son_hata:
        print("Jarvis: Arama motorları sonuç vermedi:", son_hata)

    return []


def web_arastir(sorgu):
    print("Jarvis: Araştırıyorum...")

    sonuclar = ddgs_yedekli_ara(sorgu, 3)

    if not sonuclar:
        print("Jarvis: Arama sonucu bulunamadı.")
        return

    kaynaklar = []

    for i, x in enumerate(sonuclar, 1):
        kaynaklar.append(
            f"{i}. BAŞLIK: {x.get('title', '')}\n"
            f"ÖZET: {x.get('body', '')}\n"
            f"URL: {x.get('href', '')}"
        )

    prompt = (
        "Aşağıdaki web arama sonuçlarını kullanarak Türkçe, kısa ve "
        "doğrulanabilir bir özet hazırla. Kaynaklarda olmayan bilgiyi "
        "uydurma. En fazla 6 kısa madde kullan.\n\n"
        f"SORU: {sorgu}\n\n"
        + "\n\n".join(kaynaklar)
    )

    try:
        r = subprocess.run(
            ["cmd", "/c", "claude", "-p"],
            input=prompt,
            capture_output=True,
            text=True,
            encoding="utf-8",
            timeout=180
        )

        if r.returncode == 0 and r.stdout.strip():
            print("\nJarvis araştırma özeti:\n")
            print(r.stdout.strip())
        else:
            print("\n".join(kaynaklar))

    except Exception as e:
        print("Jarvis: Claude çağrısı başarısız:", e)
        print("\n".join(kaynaklar))


# ============================================================
# GÖRÜNÜR ARAŞTIRMA / BELGE GÜVENLİĞİ
# ============================================================

def belge_parmak_izi(veri, uzanti):
    uzanti = uzanti.lower()

    if uzanti == ".docx":
        try:
            with zipfile.ZipFile(io.BytesIO(veri), "r") as z:
                xml = z.read("word/document.xml")

            xml = re.sub(rb">\s+<", b"><", xml)

            return hashlib.sha256(xml).hexdigest()

        except Exception:
            pass

    return hashlib.sha256(veri).hexdigest()


def dosya_parmak_izi(yol):
    try:
        yol = Path(yol)

        return belge_parmak_izi(
            yol.read_bytes(),
            yol.suffix.lower()
        )

    except Exception:
        return None


def mevcut_belgeleri_tara():
    sonuc = {}

    if not INDIRMELER.exists():
        return sonuc

    for p in INDIRMELER.iterdir():
        if p.is_file() and p.suffix.lower() in BELGE_UZANTILARI:
            h = dosya_parmak_izi(p)

            if h:
                sonuc[h] = p

    return sonuc


def url_uzantisi(url):
    return Path(
        unquote(
            urlparse(url).path
        )
    ).suffix.lower()


def belge_dogrula(veri, uzanti):
    if not veri:
        return False

    uzanti = uzanti.lower()

    if uzanti == ".pdf":
        return veri.startswith(b"%PDF")

    if uzanti == ".docx":
        if not veri.startswith(b"PK"):
            return False

        try:
            with zipfile.ZipFile(io.BytesIO(veri), "r") as z:
                adlar = set(z.namelist())

                return (
                    "[Content_Types].xml" in adlar
                    and "word/document.xml" in adlar
                )

        except Exception:
            return False

    if uzanti == ".doc":
        return veri.startswith(
            bytes.fromhex("D0CF11E0A1B11AE1")
        )

    return False


def url_belge_indir(url):
    uzanti = url_uzantisi(url)

    if uzanti not in BELGE_UZANTILARI:
        return None, None, "HTML/uygunsuz uzantı"

    req = Request(
        url,
        headers={"User-Agent": "Mozilla/5.0"}
    )

    try:
        with urlopen(req, timeout=25) as resp:
            uzunluk = resp.headers.get("Content-Length")

            if (
                uzunluk
                and int(uzunluk) > MAKSIMUM_DOSYA_BOYUTU
            ):
                return None, None, "Dosya çok büyük"

            veri = resp.read(
                MAKSIMUM_DOSYA_BOYUTU + 1
            )

    except Exception as e:
        return None, None, f"İndirme hatası: {e}"

    if len(veri) > MAKSIMUM_DOSYA_BOYUTU:
        return None, None, "Dosya çok büyük"

    if not belge_dogrula(veri, uzanti):
        return (
            None,
            None,
            "Gerçek belge biçimi doğrulanamadı"
        )

    ad = Path(
        unquote(
            urlparse(url).path
        )
    ).name

    if not ad:
        ad = f"belge_{int(time.time())}{uzanti}"

    return veri, guvenli_ad(ad), None


def resmi_domain_bul(sorgu):
    n = norm(sorgu)

    kurallar = [
        ("bartin", "bartin.edu.tr"),
        ("yuksekogretim", "yok.gov.tr"),
        ("yok", "yok.gov.tr"),
        ("milli egitim", "meb.gov.tr"),
        ("meb", "meb.gov.tr"),
        ("avrupa konseyi", "coe.int"),
        ("council of europe", "coe.int"),
    ]

    for anahtar, domain in kurallar:
        if anahtar in n:
            return domain

    return None


def konu_puani(sonuc, sorgu, resmi_domain=None):
    baslik = norm(sonuc.get("title", ""))
    govde = norm(sonuc.get("body", ""))
    url = norm(sonuc.get("href", ""))

    kelimeler = [
        x
        for x in norm(sorgu).split()
        if len(x) >= 4
        and x not in {
            "icin",
            "ile",
            "olan",
            "gibi",
            "formu",
            "belge",
        }
    ]

    puan = 0

    puan += sum(
        2 for k in kelimeler if k in baslik
    )

    puan += sum(
        1 for k in kelimeler if k in govde
    )

    puan += sum(
        1 for k in kelimeler if k in url
    )

    if resmi_domain and resmi_domain in url:
        puan += 15

    if url_uzantisi(
        sonuc.get("href", "")
    ) in BELGE_UZANTILARI:
        puan += 5

    return puan


def belge_adaylari(sorgu, adet=25):
    resmi = resmi_domain_bul(sorgu)
    arama_sorgusu = sorgu

    if resmi:
        arama_sorgusu += f" site:{resmi}"

    tum = []

    for motor in ARAMA_MOTORLARI:
        try:
            sonuclar = list(
                DDGS().text(
                    arama_sorgusu,
                    max_results=adet,
                    backend=motor
                )
            )

            tum.extend(sonuclar)

            if tum:
                break

        except Exception:
            continue

    benzersiz = {}

    for x in tum:
        url = x.get("href", "")

        if not url:
            continue

        if url_uzantisi(url) not in BELGE_UZANTILARI:
            continue

        if resmi and resmi not in norm(url):
            continue

        benzersiz[url] = x

    return sorted(
        benzersiz.values(),
        key=lambda x: konu_puani(
            x,
            sorgu,
            resmi
        ),
        reverse=True
    )


def claude_penceresi_bul(bekle=30):
    def ara():
        try:
            pencereler = Desktop(
                backend="uia"
            ).windows()

            tam = [
                w
                for w in pencereler
                if w.window_text().strip() == "Claude"
            ]

            if tam:
                return tam[0]

            parcali = [
                w
                for w in pencereler
                if "claude" in norm(
                    w.window_text()
                )
            ]

            if parcali:
                return parcali[0]

        except Exception:
            pass

        return None

    pencere = ara()

    if pencere:
        return pencere

    try:
        subprocess.Popen([
            "explorer.exe",
            f"shell:AppsFolder\\{CLAUDE_APP_ID}"
        ])
    except Exception:
        return None

    baslangic = time.time()

    while time.time() - baslangic < bekle:
        time.sleep(1)

        pencere = ara()

        if pencere:
            return pencere

    return None


def claude_mesaj_gonder(mesaj):
    pencere = claude_penceresi_bul()

    if not pencere:
        print(
            "Jarvis: Claude Desktop penceresi "
            "bulunamadı."
        )
        return False

    try:
        try:
            pencere.restore()
        except Exception:
            pass

        pencere.set_focus()

        time.sleep(1)

        editler = [
            x
            for x in pencere.descendants(
                control_type="Edit"
            )
            if x.is_visible()
            and x.is_enabled()
        ]

        if not editler:
            print(
                "Jarvis: Claude mesaj kutusu "
                "bulunamadı."
            )
            return False

        hedef = max(
            editler,
            key=lambda x: x.rectangle().top
        )

        hedef.click_input()

        pyperclip.copy(mesaj)

        keyboard.send_keys("^v")

        time.sleep(0.4)

        keyboard.send_keys("{ENTER}")

        return True

    except Exception as e:
        print(
            "Jarvis: Claude'a mesaj "
            "gönderilemedi:",
            e
        )
        return False


def markdown_docx_cevir(md_yolu):
    md_yolu = Path(md_yolu)

    docx_yolu = benzersiz_yol(
        md_yolu.with_suffix(".docx")
    )

    doc = Document()

    try:
        metin = md_yolu.read_text(
            encoding="utf-8"
        )
    except UnicodeDecodeError:
        metin = md_yolu.read_text(
            encoding="utf-8",
            errors="replace"
        )

    for satir in metin.splitlines():
        s = satir.rstrip()

        if s.startswith("### "):
            doc.add_heading(
                s[4:],
                level=3
            )

        elif s.startswith("## "):
            doc.add_heading(
                s[3:],
                level=2
            )

        elif s.startswith("# "):
            doc.add_heading(
                s[2:],
                level=1
            )

        elif s.startswith("- "):
            doc.add_paragraph(
                s[2:],
                style="List Bullet"
            )

        elif re.match(r"^\d+\.\s+", s):
            doc.add_paragraph(
                re.sub(
                    r"^\d+\.\s+",
                    "",
                    s
                ),
                style="List Number"
            )

        elif not s:
            doc.add_paragraph("")

        else:
            doc.add_paragraph(s)

    doc.save(docx_yolu)

    return docx_yolu


def gorunur_arastir(sorgu):
    print(
        "Jarvis: Görünür araştırma başlıyor."
    )

    google_ara(sorgu)

    time.sleep(2)

    adaylar = belge_adaylari(
        sorgu,
        30
    )

    if not adaylar:
        print(
            "Jarvis: Güvenli ve doğrudan "
            "indirilebilir PDF/DOCX/DOC "
            "bulamadım."
        )
        return

    mevcut = mevcut_belgeleri_tara()
    secilen_hashler = set()
    indirilenler = []

    for sonuc in adaylar:
        if len(indirilenler) >= MAKSIMUM_BELGE:
            break

        url = sonuc.get("href", "")
        baslik = (
            sonuc.get("title", "")
            or url
        )

        print("\nJarvis: Aday:", baslik)
        print("Jarvis:", url)

        tarayicida_ac(url)

        time.sleep(1)

        veri, ad, hata = url_belge_indir(
            url
        )

        if hata:
            print(
                "Jarvis: RED ->",
                hata
            )
            continue

        h = belge_parmak_izi(
            veri,
            Path(ad).suffix.lower()
        )

        if h in secilen_hashler:
            print(
                "Jarvis: AYNI BELGE -> "
                "Bu içerik bu araştırmada "
                "daha önce bulundu. "
                "Tekrar indirmiyorum."
            )
            continue

        if h in mevcut:
            p = mevcut[h]

            print(
                "Jarvis: AYNI BELGE -> "
                "Bu belge bilgisayarda "
                "zaten var:",
                p.name
            )

            print(
                "Jarvis: Mevcut kopyayı "
                "kullanıyorum."
            )

            secilen_hashler.add(h)
            indirilenler.append(p)
            continue

        hedef = benzersiz_yol(
            INDIRMELER / ad
        )

        try:
            hedef.write_bytes(veri)

            secilen_hashler.add(h)

            mevcut[h] = hedef

            indirilenler.append(hedef)

            print(
                "Jarvis: İNDİRİLDİ ->",
                hedef.name
            )

        except Exception as e:
            print(
                "Jarvis: Kaydedilemedi:",
                e
            )

    if not indirilenler:
        print(
            "Jarvis: Kullanılabilir "
            "benzersiz belge indirilemedi."
        )
        return

    klasoru_ac(INDIRMELER)

    time.sleep(2)

    damga = time.strftime(
        "%Y%m%d_%H%M%S"
    )

    md_adi = (
        f"Jarvis_Arastirma_{damga}.md"
    )

    md_yolu = (
        INDIRMELER / md_adi
    )

    dosya_listesi = "\n".join(
        f"- {p.name}"
        for p in indirilenler
    )

    mesaj = (
        "Downloads klasöründeki aşağıdaki "
        "belgeleri incele:\n"
        f"{dosya_listesi}\n\n"
        f"Araştırma sorusu: {sorgu}\n\n"
        "Yalnızca bu belgelerin desteklediği "
        "bilgilerle Türkçe bir araştırma "
        "özeti hazırla. Kaynakta olmayan "
        "bilgiyi uydurma. Önemli bulguları "
        "ve mümkünse belge adlarını belirt. "
        "Sonucu Downloads klasörüne tam "
        "olarak şu adla Markdown dosyası "
        f"olarak kaydet: {md_adi}"
    )

    print(
        "Jarvis: Claude Desktop'a "
        "belgeler gönderiliyor..."
    )

    if not claude_mesaj_gonder(mesaj):
        return

    print(
        "Jarvis: Claude'un Markdown "
        "dosyasını oluşturmasını "
        "bekliyorum..."
    )

    baslangic = time.time()

    while time.time() - baslangic < 180:
        if (
            md_yolu.exists()
            and md_yolu.stat().st_size > 0
        ):
            print(
                "Jarvis: Araştırma dosyası "
                "oluştu:",
                md_yolu.name
            )

            docx = markdown_docx_cevir(
                md_yolu
            )

            print(
                "Jarvis: Word dosyası "
                "oluşturuldu:",
                docx.name
            )

            dosyayi_ac(docx)

            return

        time.sleep(3)

    print(
        "Jarvis: 3 dakika içinde Markdown "
        "dosyası oluşmadı. Claude açık "
        "bırakıldı."
    )


# ============================================================
# HIGGSFIELD
# ŞİMDİLİK SADECE KREDİ HARCAMAYAN GÜVENLİ KOMUTLAR
# ============================================================

def hf_json(args, timeout=180):
    """
    Windows'ta Higgsfield npm CLI'yi cmd /c üzerinden çalıştırır.
    """

    komut = (
        ["cmd", "/c", "higgsfield"]
        + list(args)
        + ["--json"]
    )

    try:
        r = subprocess.run(
            komut,
            capture_output=True,
            text=True,
            encoding="utf-8",
            timeout=timeout
        )

    except subprocess.TimeoutExpired:
        return (
            None,
            "Higgsfield komutu zaman aşımına uğradı."
        )

    except Exception as e:
        return (
            None,
            f"Higgsfield çalıştırılamadı: {e}"
        )

    if r.returncode != 0:
        hata = (
            r.stderr.strip()
            or r.stdout.strip()
            or f"Çıkış kodu: {r.returncode}"
        )

        return None, hata

    try:
        return json.loads(r.stdout), None

    except Exception as e:
        return (
            None,
            "Higgsfield JSON çıktısı "
            f"okunamadı: {e}\n"
            f"Ham çıktı:\n{r.stdout[:1000]}"
        )


def hf_status():
    data, hata = hf_json([
        "account",
        "status"
    ])

    if hata:
        print(
            "Jarvis:",
            hata
        )
        return

    print(
        "Jarvis: Higgsfield hesabı bağlı."
    )

    if not isinstance(data, dict):
        print(
            json.dumps(
                data,
                ensure_ascii=False,
                indent=2
            )
        )
        return

    email = (
        data.get("email")
        or data.get("account_email")
    )

    user = data.get("user")

    if (
        email is None
        and isinstance(user, dict)
    ):
        email = user.get("email")

    plan = (
        data.get("plan")
        or data.get("plan_name")
    )

    workspace = data.get("workspace")

    if isinstance(workspace, dict):
        if plan is None:
            plan = workspace.get("plan")

    credits = (
        data.get("credits")
        or data.get("available_credits")
    )

    if isinstance(workspace, dict):
        if credits is None:
            credits = workspace.get("credits")

    if email is not None:
        print(
            "Jarvis: E-posta:",
            email
        )

    if plan is not None:
        print(
            "Jarvis: Plan:",
            plan
        )

    if credits is not None:
        print(
            "Jarvis: Kredi:",
            credits
        )

    if (
        email is None
        and plan is None
        and credits is None
    ):
        print(
            json.dumps(
                data,
                ensure_ascii=False,
                indent=2
            )
        )


def hf_cost_test():
    data, hata = hf_json([
        "generate",
        "cost",
        "seedance_2_5",
        "--prompt",
        "test",
        "--duration",
        "5",
        "--resolution",
        "480p",
    ])

    if hata:
        print(
            "Jarvis:",
            hata
        )
        return

    if (
        isinstance(data, dict)
        and "credits" in data
    ):
        print(
            "Jarvis: Tahmini maliyet:",
            data["credits"],
            "kredi"
        )

    else:
        print(
            "Jarvis: Maliyet çıktısı:"
        )

        print(
            json.dumps(
                data,
                ensure_ascii=False,
                indent=2
            )
        )


def hf_son_tamamlanmis_seedance():
    data, hata = hf_json([
        "generate",
        "list"
    ])

    if hata:
        return None, hata

    if not isinstance(data, list):
        return (
            None,
            "Higgsfield iş listesi "
            "beklenen biçimde gelmedi."
        )

    for x in data:
        if (
            isinstance(x, dict)
            and x.get("job_type")
            == "seedance_2_5"
            and x.get("status")
            == "completed"
            and x.get("result_url")
        ):
            return x, None

    return (
        None,
        "Tamamlanmış Seedance 2.5 "
        "videosu bulunamadı."
    )


def hf_son_video_indir():
    job, hata = (
        hf_son_tamamlanmis_seedance()
    )

    if hata:
        print(
            "Jarvis:",
            hata
        )
        return

    url = job["result_url"]

    job_id = job.get(
        "id",
        "bilinmiyor"
    )

    hedef = benzersiz_yol(
        INDIRMELER
        / "Jarvis_Higgsfield_Test.mp4"
    )

    try:
        req = Request(
            url,
            headers={
                "User-Agent": "Mozilla/5.0"
            }
        )

        with (
            urlopen(
                req,
                timeout=60
            ) as resp,
            open(
                hedef,
                "wb"
            ) as f
        ):
            shutil.copyfileobj(
                resp,
                f
            )

        print(
            "Jarvis: Job:",
            job_id
        )

        print(
            "Jarvis: Video indirildi:",
            hedef
        )

    except Exception as e:
        print(
            "Jarvis: Video indirilemedi:",
            e
        )



def hf_json_icinden_deger_bul(veri, anahtarlar):
    if isinstance(veri, dict):
        for anahtar in anahtarlar:
            deger = veri.get(anahtar)

            if deger not in (None, "", []):
                return deger

        for deger in veri.values():
            bulunan = hf_json_icinden_deger_bul(
                deger,
                anahtarlar
            )

            if bulunan not in (None, "", []):
                return bulunan

    elif isinstance(veri, list):
        for oge in veri:
            bulunan = hf_json_icinden_deger_bul(
                oge,
                anahtarlar
            )

            if bulunan not in (None, "", []):
                return bulunan

    return None


def hf_video_url_bul(veri):
    if isinstance(veri, dict):
        for anahtar in (
            "result_url",
            "output_url",
            "video_url",
            "min_result_url",
        ):
            deger = veri.get(anahtar)

            if (
                isinstance(deger, str)
                and deger.startswith(
                    ("http://", "https://")
                )
            ):
                return deger

        for deger in veri.values():
            bulunan = hf_video_url_bul(deger)

            if bulunan:
                return bulunan

    elif isinstance(veri, list):
        for oge in veri:
            bulunan = hf_video_url_bul(oge)

            if bulunan:
                return bulunan

    elif isinstance(veri, str):
        alt = veri.lower()

        if (
            veri.startswith(
                ("http://", "https://")
            )
            and (
                ".mp4" in alt
                or ".mov" in alt
                or ".webm" in alt
            )
        ):
            return veri

    return None


def hf_video_url_indir(url):
    try:
        url_adi = Path(
            unquote(
                urlparse(url).path
            )
        ).name

        if not url_adi:
            url_adi = (
                "Higgsfield_Seedance_"
                + time.strftime(
                    "%Y%m%d_%H%M%S"
                )
                + ".mp4"
            )

        if not Path(url_adi).suffix:
            url_adi += ".mp4"

        hedef = benzersiz_yol(
            INDIRMELER
            / guvenli_ad(url_adi)
        )

        req = Request(
            url,
            headers={
                "User-Agent": "Mozilla/5.0"
            }
        )

        with (
            urlopen(
                req,
                timeout=120
            ) as resp,
            open(
                hedef,
                "wb"
            ) as f
        ):
            shutil.copyfileobj(
                resp,
                f
            )

        print(
            "Jarvis: Video indirildi:",
            hedef
        )

        return hedef

    except Exception as e:
        print(
            "Jarvis: Video indirilemedi:",
            e
        )

        return None


def hf_video_uret():
    print()

    print(
        "Jarvis: Seedance 2.5 "
        "video üretim modu."
    )

    print(
        "Jarvis: Çözünürlük: 480p"
    )

    print(
        "Jarvis: En-boy oranı: 16:9"
    )

    print(
        "Jarvis: Şimdilik referans medya "
        "olmadan T2V üretimi kullanılıyor."
    )

    prompt = input(
        "Jarvis: Promptu yaz:\n> "
    ).strip()

    if not prompt:
        print(
            "Jarvis: Prompt boş. "
            "İşlem iptal edildi."
        )

        return

    sure_raw = input(
        "Jarvis: Süre kaç saniye? "
        "(boş bırakırsan 5): "
    ).strip()

    if not sure_raw:
        sure = 5

    else:
        try:
            sure = int(sure_raw)

        except ValueError:
            print(
                "Jarvis: Süre tam sayı olmalı. "
                "İşlem iptal edildi."
            )

            return

    if sure <= 0:
        print(
            "Jarvis: Süre sıfırdan "
            "büyük olmalı."
        )

        return

    parametreler = [
        "--prompt",
        prompt,
        "--duration",
        str(sure),
        "--resolution",
        "480p",
    ]

    print()

    print(
        "Jarvis: Önce maliyeti "
        "hesaplıyorum. "
        "Bu adım kredi harcamaz."
    )

    maliyet_data, hata = hf_json([
        "generate",
        "cost",
        "seedance_2_5",
        *parametreler,
    ])

    if hata:
        print(
            "Jarvis: Maliyet "
            "hesaplanamadı:"
        )

        print(hata)

        return

    kredi = (
        maliyet_data.get("credits")
        if isinstance(
            maliyet_data,
            dict
        )
        else None
    )

    if kredi is None:
        print(
            "Jarvis: Maliyet çıktısını "
            "okuyamadım."
        )

        print(
            json.dumps(
                maliyet_data,
                ensure_ascii=False,
                indent=2
            )
        )

        return

    print()

    print(
        "Jarvis: Model: Seedance 2.5"
    )

    print(
        "Jarvis: Süre:",
        sure,
        "saniye"
    )

    print(
        "Jarvis: Çözünürlük: 480p"
    )

    print(
        "Jarvis: En-boy oranı: 16:9"
    )

    print(
        "Jarvis: Tahmini maliyet:",
        kredi,
        "kredi"
    )

    onay = norm(
        input(
            "Jarvis: Üretimi başlatayım mı? "
            "(EVET/HAYIR): "
        )
    )

    if onay != "evet":
        print(
            "Jarvis: İptal edildi. "
            "Video üretimi başlatılmadı."
        )

        return

    print()

    print(
        "Jarvis: ONAY ALINDI."
    )

    print(
        "Jarvis: Kredi harcayan üretim "
        "işi başlatılıyor..."
    )

    uretim_data, hata = hf_json([
        "generate",
        "create",
        "seedance_2_5",
        *parametreler,
        "--wait",
        "--wait-timeout",
        "20m",
        "--wait-interval",
        "5s",
    ], timeout=1500)

    if hata:
        print(
            "Jarvis: Higgsfield "
            "üretim hatası:"
        )

        print(hata)

        return

    job_id = hf_json_icinden_deger_bul(
        uretim_data,
        (
            "id",
            "job_id"
        )
    )

    video_url = hf_video_url_bul(
        uretim_data
    )

    if job_id:
        print(
            "Jarvis: Job:",
            job_id
        )

    if (
        not video_url
        and job_id
    ):
        print(
            "Jarvis: Sonuç URL'si için "
            "işi tekrar kontrol ediyorum..."
        )

        bekleme_data, bekleme_hata = hf_json([
            "generate",
            "wait",
            str(job_id),
            "--timeout",
            "20m",
            "--interval",
            "5s",
            "--quiet",
        ], timeout=1500)

        if not bekleme_hata:
            video_url = hf_video_url_bul(
                bekleme_data
            )

    if (
        not video_url
        and job_id
    ):
        liste_data, liste_hata = hf_json([
            "generate",
            "list"
        ])

        if (
            not liste_hata
            and isinstance(
                liste_data,
                list
            )
        ):
            for x in liste_data:
                if (
                    isinstance(x, dict)
                    and str(
                        x.get("id")
                    ) == str(job_id)
                    and x.get("status")
                    == "completed"
                    and x.get(
                        "result_url"
                    )
                ):
                    video_url = x[
                        "result_url"
                    ]

                    break

    if not video_url:
        print(
            "Jarvis: Üretim tamamlanmış "
            "olabilir fakat sonuç URL'sini "
            "otomatik bulamadım."
        )

        if job_id:
            print(
                "Jarvis: Job ID:",
                job_id
            )

        return

    print(
        "Jarvis: Video tamamlandı."
    )

    print(
        "Jarvis: Sonuç indiriliyor..."
    )

    hedef = hf_video_url_indir(
        video_url
    )

    if hedef:
        print(
            "Jarvis: Higgsfield üretim "
            "akışı tamamlandı."
        )


# ============================================================


def godot_test_projesini_ac():
    from pathlib import Path
    import subprocess

    godot_exe = Path(
        r"C:\Users\Halit Fatih Böcek\Documents\Godot_4.7.1\Godot_v4.7.1-stable_win64.exe"
    )

    proje = Path(
        r"C:\Users\Halit Fatih Böcek\Documents\Cemil_Jarvis_Test"
    )

    project_file = proje / "project.godot"

    if not godot_exe.exists():
        print("Jarvis: Godot bulunamadı:")
        print(godot_exe)
        return

    if not project_file.exists():
        print("Jarvis: Cemil test projesi bulunamadı:")
        print(project_file)
        return

    print("Jarvis: Cemil Godot TEST projesi açılıyor.")
    print("Jarvis: Ana projeye dokunulmuyor.")
    print(f"Jarvis: Proje: {proje}")

    try:
        subprocess.Popen(
            [
                str(godot_exe),
                "--editor",
                "--path",
                str(proje),
            ],
            cwd=str(proje),
        )
        print("Jarvis: Godot test projesi başlatıldı.")
    except Exception as e:
        print(f"Jarvis: Godot açılamadı: {e}")


# KOMUTLAR
# ============================================================


def ai_dogal_dil_fallback(komut):
    """
    Mevcut Jarvis kurallari komutu anlayamazsa Claude'a yalnizca
    kullanicinin niyetini cozdurur.

    Claude bilgisayarda dogrudan islem yapmaz.
    Islemi Jarvis'in izin verilen fonksiyonlari gerceklestirir.
    """

    komut = str(komut).strip()

    if not komut:
        return True

    prompt = f"""
Sen Windows'ta calisan Jarvis isimli yerel asistana ait
SADECE dogal dil niyet yorumlayicisisin.

Kullanicinin Turkce cumlesini anla.

SADECE gecerli JSON dondur.
Markdown, aciklama veya kod blogu kullanma.

Izin verilen eylemler:

ac
kapat
ara
arastir
bul
sil
godot_test_ac
higgsfield_durum
higgsfield_maliyet
higgsfield_video_uret
sohbet
bilinmiyor

JSON bicimi:

{{
  "eylem": "eylem",
  "hedef": "temizlenmis hedef veya bos",
  "yanit": "yalnizca sohbet ise kisa Turkce cevap, digerlerinde bos"
}}

Kurallar:

1. "Merhaba Jarvis bana kisisel bilgi kayit formunu acar misin"
   gibi cumle:
   eylem = ac
   hedef = kisisel bilgi kayit formu

2. Hedeften Turkce ekleri ve gereksiz nezaket kelimelerini temizle.
   Ornek:
   "tez dosyami acar misin" -> hedef "tez"
   "Chrome'u acar misin" -> hedef "chrome"

3. Program, dosya veya klasor acma taleplerinin hepsi "ac" olsun.

4. Bir programi kapatma istegi "kapat" olsun.

5. Google/internet aramasi "ara" olsun.

6. Konu hakkinda bilgi toplayip ozetleme istegi "arastir" olsun.

7. Dosya veya klasoru sadece bulma istegi "bul" olsun.

8. Silme istegi "sil" olsun.
   Silme islemi daha sonra Jarvis tarafinda tekrar onaylanacaktir.

9. Cemil/Godot test projesini acma istegi "godot_test_ac" olsun.

10. Higgsfield hesap/durum sorgusu "higgsfield_durum" olsun.

11. Higgsfield maliyet sorgusu "higgsfield_maliyet" olsun.

12. Higgsfield video uretme istegi "higgsfield_video_uret" olsun.
    Kredi harcamasi Jarvis tarafinda ayrica onaylanacaktir.

13. Kullanici sadece konusuyor, soru soruyor veya sohbet ediyorsa
    "sohbet" kullan ve "yanit" alaninda kisa, dogal Turkce cevap ver.

14. Emin degilsen "bilinmiyor" kullan. Tahmin etme.

KULLANICI CUMLESI:
{komut}
"""

    try:
        sonuc = subprocess.run(
            [
                "cmd",
                "/c",
                "claude",
                "-p"
            ],
            input=prompt,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=60
        )

    except subprocess.TimeoutExpired:
        print(
            "Jarvis: Dogal dil yorumlama zaman asimina ugradi."
        )
        return True

    except Exception as e:
        print(
            "Jarvis: Claude dogal dil hatasi:",
            e
        )
        return True

    if sonuc.returncode != 0:
        hata = (
            sonuc.stderr.strip()
            or sonuc.stdout.strip()
        )

        print(
            "Jarvis: Claude dogal dil yorumlayicisi calismadi:",
            hata
        )
        return True

    ham = sonuc.stdout.strip()

    if not ham:
        print(
            "Jarvis: Komutu yorumlayamadim."
        )
        return True

    # Claude nadiren ```json ... ``` dondururse temizle.
    ham = re.sub(
        r"^```(?:json)?\s*",
        "",
        ham,
        flags=re.I
    )

    ham = re.sub(
        r"\s*```$",
        "",
        ham
    ).strip()

    # JSON disinda fazladan metin olursa yalnizca ilk nesneyi al.
    eslesme = re.search(
        r"\{.*\}",
        ham,
        flags=re.S
    )

    if eslesme:
        ham = eslesme.group(0)

    try:
        veri = json.loads(
            ham
        )

    except Exception:
        print(
            "Jarvis: AI yanitini anlayamadim."
        )
        return True

    eylem = norm(
        veri.get(
            "eylem",
            "bilinmiyor"
        )
    )

    hedef_raw = str(
        veri.get(
            "hedef",
            ""
        )
    ).strip()

    yanit = str(
        veri.get(
            "yanit",
            ""
        )
    ).strip()

    # --------------------------------------------------------
    # AC
    # --------------------------------------------------------

    if eylem == "ac":

        if not hedef_raw:
            print(
                "Jarvis: Neyi acmami istedigini anlayamadim."
            )
            konus(
                "Neyi açmamı istediğini anlayamadım."
            )
            return True

        program = program_adi_bul(
            hedef_raw
        )

        if program:
            program_ac(
                program
            )
            return True

        return dogal_ac_komutu(
            hedef_raw
        )

    # --------------------------------------------------------
    # PROGRAM KAPAT
    # --------------------------------------------------------

    if eylem == "kapat":

        program = program_adi_bul(
            hedef_raw
        )

        if program:
            program_kapat(
                program
            )
        else:
            print(
                "Jarvis: Kapatilacak programi bulamadim."
            )
            konus(
                "Kapatılacak programı bulamadım."
            )

        return True

    # --------------------------------------------------------
    # GOOGLE ARA
    # --------------------------------------------------------

    if eylem == "ara":

        if hedef_raw:
            google_ara(
                hedef_raw
            )
        else:
            konus(
                "Ne aramamı istediğini anlayamadım."
            )

        return True

    # --------------------------------------------------------
    # ARASTIR
    # --------------------------------------------------------

    if eylem == "arastir":

        if hedef_raw:
            web_arastir(
                hedef_raw
            )
        else:
            konus(
                "Neyi araştırmamı istediğini anlayamadım."
            )

        return True

    # --------------------------------------------------------
    # DOSYA / KLASOR BUL
    # --------------------------------------------------------

    if eylem == "bul":

        if not hedef_raw:
            konus(
                "Neyi bulmamı istediğini anlayamadım."
            )
            return True

        dosyalar = dosya_bul(
            hedef_raw,
            en_fazla=5
        )

        klasorler = klasor_bul(
            hedef_raw,
            en_fazla=5
        )

        if not dosyalar and not klasorler:
            print(
                f"Jarvis: '{hedef_raw}' bulunamadi."
            )
            konus(
                "Aradığın dosya veya klasörü bulamadım."
            )
            return True

        print(
            "Jarvis: Bulduklarim:"
        )

        for p in dosyalar:
            print(
                " -",
                p
            )

        for p in klasorler:
            print(
                " -",
                p
            )

        adet = len(
            dosyalar
        ) + len(
            klasorler
        )

        konus(
            f"{adet} eşleşme buldum."
        )

        return True

    # --------------------------------------------------------
    # SIL
    # --------------------------------------------------------

    if eylem == "sil":

        if hedef_raw:
            # Mevcut fonksiyon EVET/HAYIR onayi ister.
            dosya_sil_komutu(
                hedef_raw
            )
        else:
            konus(
                "Hangi dosyayı silmemi istediğini anlayamadım."
            )

        return True

    # --------------------------------------------------------
    # GODOT TEST
    # --------------------------------------------------------

    if eylem == "godot_test_ac":
        godot_test_projesini_ac()
        return True

    # --------------------------------------------------------
    # HIGGSFIELD
    # --------------------------------------------------------

    if eylem == "higgsfield_durum":
        hf_status()
        return True

    if eylem == "higgsfield_maliyet":
        hf_cost_test()
        return True

    if eylem == "higgsfield_video_uret":
        # Mevcut hf_video_uret fonksiyonunda kredi onayi korunur.
        hf_video_uret()
        return True

    # --------------------------------------------------------
    # SOHBET
    # --------------------------------------------------------

    if eylem == "sohbet":

        if not yanit:
            yanit = (
                "Seni anladım ama buna verecek "
                "uygun bir cevap oluşturamadım."
            )

        print(
            "Jarvis:",
            yanit
        )

        konus(
            yanit
        )

        return True

    print(
        "Jarvis: Ne yapmak istedigini tam olarak anlayamadim."
    )

    konus(
        "Ne yapmak istediğini tam olarak anlayamadım."
    )

    return True




def komutu_isle_ai_first(komut):
    """
    Sesli Jarvis icin ana karar katmani.

    Kullanici dogal Turkce konusur.
    Claude niyeti yorumlar.
    Gercek Windows islemlerini Jarvis yapar.

    Eski regex sistemi silinmez; ancak ana sesli akista
    karar verici olarak kullanilmaz.
    """

    k = norm(komut)

    if not k:
        return True

    # Cikis komutlarini Claude'a gonderme.
    if k in {
        "cik",
        "kapat",
        "jarvis kapat",
        "jarvisi kapat",
        "jarvis cik",
        "exit",
        "quit",
    }:
        print(
            "Jarvis: Görüşürüz."
        )

        konus(
            "Görüşürüz."
        )

        return False

    print(
        "Jarvis: İsteğini anlıyorum..."
    )

    return ai_dogal_dil_fallback(
        komut
    )



def komutu_isle(komut):
    k = norm(komut)

    if not k:
        return True

    if k in {
        "cik",
        "kapat",
        "jarvis kapat",
        "exit",
        "quit",
    }:
        print(
            "Jarvis: Görüşürüz."
        )
        return False

    # --------------------------------------------------------
    # Higgsfield güvenli komutları
    # --------------------------------------------------------

    if k in {
        "higgsfield durumunu goster",
        "higgsfield durumu",
        "higgsfield hesap durumu",
        "higgsfield hesabini goster",
    }:
        hf_status()
        return True

    if k in {
        "higgsfield maliyet test",
        "higgsfield maliyet testi",
        "higgsfield maliyetini test et",
    }:
        hf_cost_test()
        return True

    if k in {
        "son higgsfield videosunu indir",
        "higgsfield son videoyu indir",
        "son seedance videosunu indir",
    }:
        hf_son_video_indir()
        return True


    if k in {
        "higgsfield video uret",
        "seedance video uret",
        "video uret higgsfield",
    }:
        hf_video_uret()
        return True

    # --------------------------------------------------------

    if k in {
        "godot test projesini ac",
        "cemil test projesini ac",
        "test godot projesini ac",
    }:
        godot_test_projesini_ac()
        return True

    # Masaüstü
    # --------------------------------------------------------

    if k in {
        "masaustunu ac",
        "masaustu ac",
        "masaustunu goster",
    }:
        klasoru_ac(
            MASAUSTU
        )

        print(
            "Jarvis: Masaüstü açıldı."
        )

        return True

    # --------------------------------------------------------
    # Google
    # --------------------------------------------------------

    m = re.match(
        r"^(?:google'?da|internette)\s+(.+?)\s+ara$",
        k
    )

    if m:
        sorgu = re.sub(
            r"(?i)^\s*(?:google['’]?da|internette)\s+",
            "",
            komut
        )

        sorgu = re.sub(
            r"(?i)\s+ara\s*$",
            "",
            sorgu
        ).strip()

        google_ara(
            sorgu
        )

        return True

    # --------------------------------------------------------
    # Görünür araştırma
    # --------------------------------------------------------

    if k.startswith(
        "gorunur arastir "
    ):
        sorgu = re.sub(
            r"(?i)^\s*g[öo]r[üu]n[üu]r\s+"
            r"ara[sş]t[ıi]r\s+",
            "",
            komut
        ).strip()

        gorunur_arastir(
            sorgu
        )

        return True

    # --------------------------------------------------------
    # Hızlı araştırma
    # --------------------------------------------------------

    if k.startswith(
        "arastir "
    ):
        sorgu = re.sub(
            r"(?i)^\s*ara[sş]t[ıi]r\s+",
            "",
            komut
        ).strip()

        web_arastir(
            sorgu
        )

        return True

    # --------------------------------------------------------
    # Program aç / kapat
    # --------------------------------------------------------

    if (
        k.endswith(" ac")
        or k.endswith(" acabilir misin")
    ):
        ad = program_adi_bul(
            k
        )

        if ad:
            program_ac(
                ad
            )

            return True

    if k.endswith(
        " kapat"
    ):
        ad = program_adi_bul(
            k
        )

        if ad:
            program_kapat(
                ad
            )

            return True

    # --------------------------------------------------------
    # Konum özel yeniden adlandırma
    # --------------------------------------------------------

    m = re.match(
        r"^(.+?) klasorundeki "
        r"(.+?) dosyasinin adini "
        r"(.+?) yap$",
        k
    )

    if m:
        klasor_adi = m.group(1)
        dosya_adi = m.group(2)
        yeni_ad_norm = m.group(3)

        klasor = tek_klasor(
            klasor_adi
        )

        if klasor:
            raw = re.split(
                r"(?i)\s+dosyas[ıi]n[ıi]n\s+"
                r"ad[ıi]n[ıi]\s+",
                komut,
                maxsplit=1
            )

            yeni_ad = (
                raw[1]
                if len(raw) == 2
                else yeni_ad_norm
            )

            yeni_ad = re.sub(
                r"(?i)\s+yap\s*$",
                "",
                yeni_ad
            ).strip()

            dosya_yeniden_adlandir(
                dosya_adi,
                yeni_ad,
                kok=klasor
            )

        return True

    # --------------------------------------------------------
    # Genel yeniden adlandırma
    # --------------------------------------------------------

    m = re.match(
        r"^(.+?) dosyasinin adini "
        r"(.+?) yap$",
        k
    )

    if m:
        eski = m.group(1)
        yeni_norm = m.group(2)

        raw = re.split(
            r"(?i)\s+dosyas[ıi]n[ıi]n\s+"
            r"ad[ıi]n[ıi]\s+",
            komut,
            maxsplit=1
        )

        yeni = (
            raw[1]
            if len(raw) == 2
            else yeni_norm
        )

        yeni = re.sub(
            r"(?i)\s+yap\s*$",
            "",
            yeni
        ).strip()

        dosya_yeniden_adlandir(
            eski,
            yeni
        )

        return True

    # --------------------------------------------------------
    # Taşı
    # --------------------------------------------------------

    m = re.match(
        r"^(.+?) dosyasini "
        r"(.+?) klasorune tasi$",
        k
    )

    if m:
        dosya_tasi_veya_kopyala(
            m.group(1),
            m.group(2),
            kopyala=False
        )

        return True

    # --------------------------------------------------------
    # Kopyala
    # --------------------------------------------------------

    m = re.match(
        r"^(.+?) dosyasini "
        r"(.+?) klasorune kopyala$",
        k
    )

    if m:
        dosya_tasi_veya_kopyala(
            m.group(1),
            m.group(2),
            kopyala=True
        )

        return True

    # --------------------------------------------------------
    # Güvenli sil
    # --------------------------------------------------------

    m = re.match(
        r"^(.+?) dosyasini sil$",
        k
    )

    if m:
        dosya_sil_komutu(
            m.group(1)
        )

        return True

    # --------------------------------------------------------
    # Dosya aç
    # --------------------------------------------------------

    m = re.match(
        r"^(.+?) dosyasini ac$",
        k
    )

    if m:
        dosya_ac_komutu(
            m.group(1)
        )

        return True

    # --------------------------------------------------------
    # Klasör aç
    # --------------------------------------------------------

    m = re.match(
        r"^(.+?) klasorunu ac$",
        k
    )

    if m:
        klasor_ac_komutu(
            m.group(1)
        )

        return True

    # --------------------------------------------------------
    # Dogal "X ac" komutu
    # Ornek: "kisisel bilgi kayit formu ac"
    # Program komutlari yukarida zaten kontrol edilmistir.
    # --------------------------------------------------------

    m = re.match(
        r"^(.+?) ac$",
        k
    )

    if m:
        dogal_ac_komutu(
            m.group(1)
        )

        return True

    # --------------------------------------------------------
    # Dosya bul
    # --------------------------------------------------------

    if k.endswith(
        " dosyasini bul"
    ):
        sorgu = k[
            :-len(" dosyasini bul")
        ].strip()

        bulunan = dosya_bul(
            sorgu
        )

        if not bulunan:
            print(
                "Jarvis: Dosya bulunamadı."
            )
        else:
            print(
                "Jarvis: Bulduklarım:"
            )

            for p in bulunan[:10]:
                print(
                    " -",
                    p
                )

        return True

    # --------------------------------------------------------
    # Klasör bul
    # --------------------------------------------------------

    if k.endswith(
        " klasorunu bul"
    ):
        sorgu = k[
            :-len(" klasorunu bul")
        ].strip()

        bulunan = klasor_bul(
            sorgu
        )

        if not bulunan:
            print(
                "Jarvis: Klasör bulunamadı."
            )
        else:
            print(
                "Jarvis: Bulduklarım:"
            )

            for p in bulunan[:10]:
                print(
                    " -",
                    p
                )

        return True

    return ai_dogal_dil_fallback(
        komut
    )


# ============================================================
# ANA DÖNGÜ
# ============================================================

def main():
    konus("Jarvis hazır.")
    print(
        "=" * 58
    )

    print(
        "JARVIS hazır."
    )

    print(
        "Sesli AI-first doğal dil modu aktif. "
        "Yazili komut yedegi de kullanilabilir."
    )

    print(
        "Higgsfield güvenli bağlantısı aktif:"
    )

    print(
        " - higgsfield durumunu göster"
    )

    print(
        " - higgsfield maliyet test"
    )

    print(
        " - son higgsfield videosunu indir"
    )

    print(
        "Çıkış: jarvis kapat"
    )

    print(
        "=" * 58
    )

    while True:
        try:
            komut = sesli_komut_al()

            if not komut:
                secim = input(
                    "Yazmak icin komutu gir veya tekrar dinlemek icin Enter: "
                ).strip()

                if not secim:
                    continue

                komut = secim

        except (
            EOFError,
            KeyboardInterrupt
        ):
            print(
                "\nJarvis: Görüşürüz."
            )

            break

        try:
            devam = komutu_isle_ai_first(
                komut
            )

            if not devam:
                break

        except Exception as e:
            print(
                "Jarvis: Beklenmeyen hata:",
                e
            )


if __name__ == "__main__":
    main()
